#!/usr/bin/env python3
"""Inspect a DOCX template without changing its styles or content."""

from __future__ import annotations

import argparse
from collections import Counter
import json
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE


def points(value) -> float | None:
    return round(value.pt, 3) if value is not None else None


def inches(value) -> float | None:
    return round(value.inches, 4) if value is not None else None


def font_info(font) -> dict[str, object]:
    rfonts = getattr(getattr(font, "_element", None), "rPr", None)
    east_asia = None
    if rfonts is not None:
        node = rfonts.rFonts
        if node is not None:
            east_asia = node.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia")
    return {
        "name": font.name,
        "eastAsia": east_asia,
        "size_pt": points(font.size),
        "bold": font.bold,
        "italic": font.italic,
        "underline": font.underline,
        "color": str(font.color.rgb) if font.color and font.color.rgb else None,
    }


def style_info(style, usage: int) -> dict[str, object]:
    info: dict[str, object] = {
        "name": style.name,
        "style_id": style.style_id,
        "type": str(style.type),
        "builtin": style.builtin,
        "usage_count": usage,
        "font": font_info(style.font),
    }
    if style.type == WD_STYLE_TYPE.PARAGRAPH:
        fmt = style.paragraph_format
        info["paragraph"] = {
            "alignment": str(fmt.alignment) if fmt.alignment is not None else None,
            "left_indent_in": inches(fmt.left_indent),
            "right_indent_in": inches(fmt.right_indent),
            "first_line_indent_in": inches(fmt.first_line_indent),
            "space_before_pt": points(fmt.space_before),
            "space_after_pt": points(fmt.space_after),
            "line_spacing": fmt.line_spacing,
            "keep_with_next": fmt.keep_with_next,
            "keep_together": fmt.keep_together,
            "page_break_before": fmt.page_break_before,
        }
        info["base_style"] = style.base_style.name if style.base_style else None
    return info


def paragraph_style_usage(document: Document) -> Counter[str]:
    counts: Counter[str] = Counter()

    def visit(paragraphs) -> None:
        for paragraph in paragraphs:
            counts[paragraph.style.name] += 1

    visit(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                visit(cell.paragraphs)
    for section in document.sections:
        visit(section.header.paragraphs)
        visit(section.footer.paragraphs)
    return counts


def profile(path: Path) -> dict[str, object]:
    document = Document(path)
    usage = paragraph_style_usage(document)
    styles = [
        style_info(style, usage.get(style.name, 0))
        for style in document.styles
        if style.type in {WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER, WD_STYLE_TYPE.TABLE}
    ]
    sections = []
    for section in document.sections:
        sections.append(
            {
                "page_width_in": inches(section.page_width),
                "page_height_in": inches(section.page_height),
                "left_margin_in": inches(section.left_margin),
                "right_margin_in": inches(section.right_margin),
                "top_margin_in": inches(section.top_margin),
                "bottom_margin_in": inches(section.bottom_margin),
                "header_distance_in": inches(section.header_distance),
                "footer_distance_in": inches(section.footer_distance),
            }
        )
    return {
        "source": str(path.resolve()),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "section_count": len(document.sections),
        "styles": styles,
        "sections": sections,
        "used_paragraph_styles": dict(usage),
    }


def markdown_report(data: dict[str, object]) -> str:
    lines = [
        f"# DOCX Style Profile",
        "",
        f"- Source: `{data['source']}`",
        f"- Paragraphs: {data['paragraph_count']}",
        f"- Tables: {data['table_count']}",
        "",
        "| Style | Type | Usage | Font | Size | Bold | Italic |",
        "|---|---|---:|---|---:|---|---|",
    ]
    for style in data["styles"]:
        font = style["font"]
        lines.append(
            f"| `{style['name']}` | {style['type']} | {style['usage_count']} | "
            f"{font['name'] or ''} | {font['size_pt'] or ''} | "
            f"{font['bold'] if font['bold'] is not None else ''} | "
            f"{font['italic'] if font['italic'] is not None else ''} |"
        )
    return "\n".join(lines) + "\n"


def main() -> int:
    parser = argparse.ArgumentParser(description="Profile styles and page settings in a DOCX template.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--json", type=Path, help="Write JSON profile to this path.")
    parser.add_argument("--markdown", type=Path, help="Write a human-readable Markdown profile.")
    args = parser.parse_args()
    if not args.docx.exists():
        parser.error(f"file not found: {args.docx}")
    data = profile(args.docx)
    output = json.dumps(data, ensure_ascii=False, indent=2)
    if args.json:
        args.json.write_text(output + "\n", encoding="utf-8")
    else:
        print(output)
    if args.markdown:
        args.markdown.write_text(markdown_report(data), encoding="utf-8")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

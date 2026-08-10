#!/usr/bin/env python3
"""Insert structured thesis content while inheriting styles from a DOCX template.

The writer never creates or rewrites styles. Every requested style must already
exist in the input template. This makes the user's Word file the formatting
authority.
"""

from __future__ import annotations

import argparse
import copy
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches

from lxml import etree

from omml import latex_to_omml, M, W


def insert_after(paragraph: Paragraph, style_name: str | None = None) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    if style_name:
        inserted.style = style_name
    return inserted


def find_anchor(document: Document, anchor: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == anchor.strip():
            return paragraph
    raise ValueError(f"anchor paragraph not found: {anchor!r}")


def style_half_points(document: Document, style_name: str | None, default: int = 24) -> int:
    if not style_name:
        return default
    size = document.styles[style_name].font.size
    return int(round(size.pt * 2)) if size else default


def add_inline(paragraph: Paragraph, item: dict[str, object], size: int) -> None:
    if "text" in item:
        paragraph.add_run(str(item["text"]))
        return
    latex = item.get("latex")
    if latex is None:
        raise ValueError("inline item must contain either 'text' or 'latex'")
    paragraph._p.append(latex_to_omml(str(latex), size=int(item.get("size", size))))


def _math_tab_run() -> etree._Element:
    """Return <m:r><w:tab/></m:r> so the tab lives inside the math object
    (native Word numbered-formula layout, no stray arrow glyphs)."""
    m_run = etree.Element(M("r"))
    w_tab = etree.SubElement(m_run, W("tab"))
    return m_run


def _fonts_sz(rpr: etree._Element, size: int) -> None:
    rfonts = etree.SubElement(rpr, W("rFonts"))
    rfonts.set(W("ascii"), "Cambria Math")
    rfonts.set(W("hAnsi"), "Cambria Math")
    sz = etree.SubElement(rpr, W("sz"))
    sz.set(W("val"), str(size))
    szcs = etree.SubElement(rpr, W("szCs"))
    szcs.set(W("val"), str(size))


def _numbered_math(latex: str, number: str | None, size: int) -> etree._Element:
    """Build the native Word/WPS numbered-formula layout (``formula # (2-1) <Enter>``).

    Word stores this as ``<m:oMathPara><m:oMath><m:eqArr>`` whose single cell
    ``<m:e>`` contains the formula, a literal ``#`` run, the number inside a
    delimiter ``<m:d>`` (bare text -- the editor renders the parentheses) and a
    trailing ``<m:ctrlPr>``.  Paragraph tab stops align formula centre / number right.
    """
    container = etree.Element(M("oMathPara"))
    math = etree.Element(M("oMath"))
    eqarr = etree.Element(M("eqArr"))

    # eqArr properties: keep maxDist (matches the user-verified native layout;
    # formulas with a single '=' are unaffected).  Cambria Math at document size.
    eqarrpr = etree.SubElement(eqarr, M("eqArrPr"))
    maxdist = etree.SubElement(eqarrpr, M("maxDist"))
    maxdist.set(M("val"), "1")
    ctrlpr = etree.SubElement(eqarrpr, M("ctrlPr"))
    wrpr = etree.SubElement(ctrlpr, W("rPr"))
    _fonts_sz(wrpr, size)

    # the single cell: formula + '#' + number delimiter + trailing ctrlPr
    e1 = etree.SubElement(eqarr, M("e"))
    formula = latex_to_omml(latex, size=size)
    for child in list(formula):
        e1.append(child)

    hash_run = etree.SubElement(e1, M("r"))
    hash_text = etree.SubElement(hash_run, M("t"))
    hash_text.text = "#"

    if number:
        num_d = etree.SubElement(e1, M("d"))
        num_dpr = etree.SubElement(num_d, M("dPr"))
        num_ctrl = etree.SubElement(num_dpr, M("ctrlPr"))
        num_rpr = etree.SubElement(num_ctrl, W("rPr"))
        _fonts_sz(num_rpr, size)
        etree.SubElement(num_rpr, W("i"))
        num_e = etree.SubElement(num_d, M("e"))
        num_run = etree.SubElement(num_e, M("r"))
        num_text = etree.SubElement(num_run, M("t"))
        num_text.text = number

    tail_ctrl = etree.SubElement(e1, M("ctrlPr"))
    tail_rpr = etree.SubElement(tail_ctrl, W("rPr"))
    _fonts_sz(tail_rpr, size)

    math.append(eqarr)
    container.append(math)
    return container


def add_display_formula(
    paragraph: Paragraph,
    latex: str,
    number: str | None,
    size: int,
    document: Document,
) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # 公式段落必须清除首行/左缩进：继承正文样式的 firstLine indent 会压缩
    # 公式可用行宽，导致长公式（如决策向量、归一化式）显示不完整/截断。
    # 只修改公式段落本身，不触碰正文段落，因此不影响正文格式。
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.left_indent = Inches(0)
    section = document.sections[0]
    width = section.page_width - section.left_margin - section.right_margin
    width_in = width.inches if hasattr(width, "inches") else float(width) / 914400
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(width_in / 2), WD_TAB_ALIGNMENT.CENTER
    )
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(width_in), WD_TAB_ALIGNMENT.RIGHT
    )
    paragraph._p.append(_numbered_math(latex, number, size))


def add_block(document: Document, previous: Paragraph, block: dict[str, object]) -> Paragraph:
    kind = str(block.get("type", "paragraph"))
    style_name = block.get("style")
    if style_name:
        style_name = str(style_name)
        document.styles[style_name]
    paragraph = insert_after(previous, style_name)
    if kind in {"heading", "paragraph"}:
        runs = block.get("runs")
        if runs is None:
            runs = [{"text": str(block.get("text", ""))}]
        for item in runs:
            add_inline(paragraph, dict(item), style_half_points(document, style_name))
    elif kind == "formula":
        add_display_formula(
            paragraph,
            str(block["latex"]),
            str(block["number"]) if block.get("number") is not None else None,
            int(block.get("size", style_half_points(document, style_name))),
            document,
        )
    else:
        raise ValueError(f"unsupported block type: {kind!r}")
    return paragraph


def write(template: Path, content_path: Path, output: Path) -> dict[str, object]:
    spec = json.loads(content_path.read_text(encoding="utf-8"))
    document = Document(template)
    anchor = find_anchor(document, str(spec["anchor"]))
    previous = anchor
    blocks = spec.get("blocks", [])
    for block in blocks:
        previous = add_block(document, previous, dict(block))
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return {
        "template": str(template.resolve()),
        "output": str(output.resolve()),
        "anchor": str(spec["anchor"]),
        "blocks_written": len(blocks),
        "style_policy": "template styles only; no style definitions changed",
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Write content using styles already present in a DOCX template.")
    parser.add_argument("template", type=Path)
    parser.add_argument("content", type=Path, help="JSON content specification")
    parser.add_argument("-o", "--output", type=Path, required=True)
    args = parser.parse_args()
    try:
        summary = write(args.template, args.content, args.output)
    except (KeyError, ValueError, TypeError) as exc:
        print(f"ERROR: {exc}")
        return 2
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())


def make_three_line_table(table) -> None:
    """Convert a python-docx Table into the standard academic three-line table.

    Standard (GB/T 7713.1 / journal editorial norms): only three horizontal rules
    -- top rule (1.5 pt), header rule (0.75 pt under the header row), bottom rule
    (1.5 pt); no vertical rules and no other horizontal rules inside the body.
    Header row text is centred.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _ALIGN
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(W("tblBorders"))
    if tblBorders is None:
        tblBorders = etree.SubElement(tblPr, W("tblBorders"))
    for edge in ("top", "bottom", "insideH", "insideV", "left", "right"):
        el = tblBorders.find(W(edge))
        if el is None:
            el = etree.SubElement(tblBorders, W(edge))
        el.set(W("val"), "none")
    nrows = len(table.rows)
    # header row: top 1.5pt + bottom 0.75pt
    for cell in table.rows[0].cells:
        _set_cell_edges(cell, top="12", bottom="6")
    # last row: bottom 1.5pt
    for cell in table.rows[nrows - 1].cells:
        _set_cell_edges(cell, bottom="12")
    # centre header text
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            p.alignment = _ALIGN.CENTER


def _set_cell_edges(cell, top: str | None = None, bottom: str | None = None) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(W("tcBorders"))
    if tcBorders is None:
        tcBorders = etree.SubElement(tcPr, W("tcBorders"))
    for edge, sz in (("top", top), ("bottom", bottom)):
        if sz is None:
            continue
        el = tcBorders.find(W(edge))
        if el is None:
            el = etree.SubElement(tcBorders, W(edge))
        el.set(W("val"), "single")
        el.set(W("sz"), sz)
        el.set(W("space"), "0")
        el.set(W("color"), "000000")
